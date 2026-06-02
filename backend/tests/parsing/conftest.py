"""Parsing test fixtures: build .xlsx/.docx fixtures programmatically."""

from __future__ import annotations

import sys
from pathlib import Path
from typing import Tuple

import pytest

BACKEND_DIR = Path(__file__).resolve().parent.parent.parent
SRC_DIR = BACKEND_DIR / "src"
for candidate in (BACKEND_DIR, SRC_DIR):
    if str(candidate) not in sys.path:
        sys.path.insert(0, str(candidate))


@pytest.fixture()
def png_factory(tmp_path: Path):
    """Return a callable that creates a small PNG and returns its path."""
    from PIL import Image

    def _make(name: str = "img.png", color: Tuple[int, int, int] = (160, 80, 120)) -> Path:
        target = tmp_path / name
        Image.new("RGB", (32, 32), color=color).save(target)
        return target

    return _make


@pytest.fixture()
def xlsx_fixture(tmp_path: Path, png_factory):
    """Build an .xlsx fixture with merged cells, hidden columns, embedded image."""
    import openpyxl
    from openpyxl.drawing.image import Image as XLImage

    img = png_factory("xfixture.png")
    xlsx_path = tmp_path / "fixture.xlsx"

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Spec"
    ws["A1"] = "Requirement ID"
    ws["B1"] = "Description"
    ws["A2"] = "REQ-001"
    ws["B2"] = "price field shall accept currency"
    ws["A4"] = "Section Header"
    ws.merge_cells("A4:B4")
    ws["B14"] = "Price: $10,450"
    ws.column_dimensions["D"].hidden = True
    image = XLImage(str(img))
    ws.add_image(image, "C14")
    wb.save(str(xlsx_path))
    return xlsx_path


@pytest.fixture()
def docx_fixture(tmp_path: Path, png_factory):
    """Build a .docx fixture with paragraphs, an inline image, and a table."""
    import docx
    from docx.shared import Inches

    img = png_factory("dfixture.png", color=(40, 200, 110))
    docx_path = tmp_path / "fixture.docx"
    d = docx.Document()
    d.add_heading("Test Cases", level=1)
    d.add_paragraph("Intro paragraph describing the test plan.")
    d.add_picture(str(img), width=Inches(0.8))
    d.add_paragraph("Paragraph after the image with details.")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "ID"
    t.cell(0, 1).text = "Description"
    t.cell(1, 0).text = "TC-001"
    t.cell(1, 1).text = "Validate price field"
    d.save(str(docx_path))
    return docx_path
